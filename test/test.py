from google.cloud import vision
import io
import os
from enum import Enum
from PIL import Image, ImageDraw
import docx

os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r'credentials.json'


def detect_document(path):
    """Detects document features in an image."""
    client = vision.ImageAnnotatorClient()

    with io.open(path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.document_text_detection(image=image)

    for page in response.full_text_annotation.pages:
        for block in page.blocks:
            print('\nBlock confidence: {}\n'.format(block.confidence))

            for paragraph in block.paragraphs:
                print('Paragraph confidence: {}'.format(
                    paragraph.confidence))

                for word in paragraph.words:
                    word_text = ''.join([
                        symbol.text for symbol in word.symbols
                    ])
                    print('Word text: {} (confidence: {})'.format(
                        word_text, word.confidence))

                    for symbol in word.symbols:
                        print('\tSymbol: {} (confidence: {})'.format(
                            symbol.text, symbol.confidence))

    if response.error.message:
        raise Exception(
            '{}\nFor more info on error messages, check: '
            'https://cloud.google.com/apis/design/errors'.format(
                response.error.message))

class FeatureType(Enum):
    PAGE = 1
    BLOCK = 2
    PARA = 3
    WORD = 4
    SYMBOL = 5

def get_document_bounds(image_file, feature):
    """Returns document bounds given an image."""
    client = vision.ImageAnnotatorClient()

    bounds = []

    with io.open(image_file, "rb") as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.document_text_detection(image=image)
    document = response.full_text_annotation

    # Collect specified feature bounds by enumerating all document features
    for page in document.pages:
        for block in page.blocks:
            for paragraph in block.paragraphs:
                for word in paragraph.words:
                    for symbol in word.symbols:
                        if feature == FeatureType.SYMBOL:
                            bounds.append(symbol.bounding_box)

                    if feature == FeatureType.WORD:
                        bounds.append(word.bounding_box)

                if feature == FeatureType.PARA:
                    bounds.append(paragraph.bounding_box)

            if feature == FeatureType.BLOCK:
                bounds.append(block.bounding_box)

    # The list `bounds` contains the coordinates of the bounding boxes.
    return bounds

def draw_boxes(image, bounds, color):
    """Draw a border around the image using the hints in the vector list."""
    draw = ImageDraw.Draw(image)

    for bound in bounds:
        draw.polygon(
            [
                bound.vertices[0].x,
                bound.vertices[0].y,
                bound.vertices[1].x,
                bound.vertices[1].y,
                bound.vertices[2].x,
                bound.vertices[2].y,
                bound.vertices[3].x,
                bound.vertices[3].y,
            ],
            None,
            color,
        )
    return image

def render_doc_text(filein, fileout):
    image = Image.open(filein)
    bounds = get_document_bounds(filein, FeatureType.BLOCK)
    draw_boxes(image, bounds, "blue")
    bounds = get_document_bounds(filein, FeatureType.PARA)
    draw_boxes(image, bounds, "red")
    bounds = get_document_bounds(filein, FeatureType.WORD)
    draw_boxes(image, bounds, "yellow")

    if fileout != 0:
        image.save(fileout)
    else:
        image.show()

def drawTextboxes(inputFile):
    fileType = (inputFile[::-1].split(".")[0])[::-1]
    outputFile = "".join(inputFile[::-1][len(fileType) + 1:])[::-1] + "_drewTextBoxes." + fileType
    render_doc_text(inputFile, outputFile)

def test_transcribeText(image_file):
    """Returns document bounds given an image."""
    fileName = image_file
    client = vision.ImageAnnotatorClient()

    bounds = []

    with io.open(image_file, "rb") as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.document_text_detection(image=image)
    document = response.full_text_annotation
    
    outDocument = docx.Document()

    for page in response.full_text_annotation.pages:
        for block in page.blocks:
            for paragraph in block.paragraphs:
                paragraph_text = ' '.join([''.join([symbol.text for symbol in word.symbols]) for word in paragraph.words])
                outDocument.add_paragraph(paragraph_text)
    
    fileType = (fileName[::-1].split(".")[0])[::-1]
    outputFile = "".join(fileName[::-1][len(fileType) + 1:])[::-1] + "_document.docx"
    outDocument.save(outputFile)

def localize_objects(path):
    """Localize objects in the local image.

    Args:
    path: The path to the local file.
    """
    from google.cloud import vision
    client = vision.ImageAnnotatorClient()

    with open(path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    objects = client.object_localization(
        image=image).localized_object_annotations

    print('Number of objects found: {}'.format(len(objects)))
    for object_ in objects:
        print('\n{} (confidence: {})'.format(object_.name, object_.score))
        print('Normalized bounding polygon vertices: ')
        for vertex in object_.bounding_poly.normalized_vertices:
            print(' - ({}, {})'.format(vertex.x, vertex.y))

def image2file(image):
    """Return `image` as PNG file-like object."""
    image_file = io.BytesIO()
    image.save(image_file, format="PNG")
    return image_file

def test_transcribeObjects(imageFile):
    """Localize objects in the local image.

    Args:
    path: The path to the local file.
    """

    img = Image.open(imageFile)
    imgheight = img.size[1]
    imgwidth = img.size[0]
    client = vision.ImageAnnotatorClient()

    outDocument = docx.Document()
    with open(imageFile, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    objects = client.object_localization(
        image=image).localized_object_annotations
    fileType = (imageFile[::-1].split(".")[0])[::-1]
    outputFile = "".join(imageFile[::-1][len(fileType) + 1:])[::-1] + "_document.docx"

    for object_ in objects:
        print(object_.bounding_poly.normalized_vertices[0].x, object_.bounding_poly.normalized_vertices[0].y)
        print(object_.bounding_poly.normalized_vertices[2].x, object_.bounding_poly.normalized_vertices[2].y)
        
        x1 = int(object_.bounding_poly.normalized_vertices[0].x*imgwidth)
        x2 = int(object_.bounding_poly.normalized_vertices[2].x*imgwidth)
        y1 = int(object_.bounding_poly.normalized_vertices[0].y*imgwidth)
        y2 = int(object_.bounding_poly.normalized_vertices[2].y*imgwidth)
        cropArea = (x1, y1, x2, y2)
        croppedImage = img.crop(cropArea)
        outDocument.add_picture(image2file(croppedImage))
    outDocument.save(outputFile)
        
def transcribeAll(imageFile):
    fileName = imageFile
    client = vision.ImageAnnotatorClient()
    bounds = []
    documentElements = {}
    img = Image.open(imageFile)
    imgheight = img.size[1]
    imgwidth = img.size[0]
    with io.open(imageFile, "rb") as image_file:
        content = image_file.read()
    image = vision.Image(content=content)
    response = client.document_text_detection(image=image)
    document = response.full_text_annotation
    outDocument = docx.Document()
    for page in response.full_text_annotation.pages:
        for block in page.blocks:
            for paragraph in block.paragraphs:
                paragraph_text = ' '.join([''.join([symbol.text for symbol in word.symbols]) for word in paragraph.words])
                tempCoords = []
                for coord in ''.join([str(vertex) for vertex in paragraph.bounding_box.vertices]).split("\n"):
                    if len(coord) > 3:
                        if coord[0] == "y":
                            tempCoords.append(int(coord[3:]))
                documentElements[paragraph_text] = max(tempCoords)
    with open(imageFile, 'rb') as image_file:
        content = image_file.read()
    objects = client.object_localization(
        image=image).localized_object_annotations
    for object_ in objects:
        x1 = int(object_.bounding_poly.normalized_vertices[0].x*imgwidth)
        x2 = int(object_.bounding_poly.normalized_vertices[2].x*imgwidth)
        y1 = int(object_.bounding_poly.normalized_vertices[0].y*imgwidth)
        y2 = int(object_.bounding_poly.normalized_vertices[2].y*imgwidth)
        cropArea = (x1, y1, x2, y2)
        croppedImage = img.crop(cropArea)
        documentElements[image2file(croppedImage)] = max(y1, y2)
    fileType = (fileName[::-1].split(".")[0])[::-1]
    outputFile = "".join(fileName[::-1][len(fileType) + 1:])[::-1] + "_document.docx"
    documentElements = dict(sorted(documentElements.items(), key=lambda item: item[1]))
    for element in documentElements.keys():
        if type(element) == str:
            outDocument.add_paragraph(element)
        else:
            outDocument.add_picture(element)
    outDocument.save(outputFile)

# drawTextboxes("test2.png")
transcribeAll("test16.jpg")